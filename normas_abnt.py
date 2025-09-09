# normas_abnt.py
# Descrição: Versão corrigida para garantir cor preta nos títulos e melhor controle de espaçamento.

from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from documento import DocumentoABNT

class MotorNormasABNT:
    def __init__(self, doc_abnt: DocumentoABNT):
        self.doc_abnt = doc_abnt
        self.MARGEM_SUPERIOR = Cm(3)
        self.MARGEM_INFERIOR = Cm(2)
        self.MARGEM_ESQUERDA = Cm(3)
        self.MARGEM_DIREITA = Cm(2)
        self.FONTE_PADRAO = 'Times New Roman'
        self.TAMANHO_FONTE_PADRAO = Pt(12)
        self.TAMANHO_FONTE_CAPA = Pt(14)
        self.COR_FONTE_PADRAO = RGBColor(0, 0, 0) # NOVO: Cor preta
        self.ESPAÇAMENTO_PADRAO = 1.5
        self.ESPAÇAMENTO_SIMPLES = 1.0
        self.RECUO_PRIMEIRA_LINHA = Cm(1.25)
        self.RECUO_CITACAO_LONGA = Cm(4)
        self.TAMANHO_FONTE_CITACAO_LONGA = Pt(10)
        self.TAMANHO_FONTE_LEGENDA = Pt(10)

    @property
    def is_artigo(self) -> bool:
        """Propriedade para verificar facilmente se o trabalho é um artigo."""
        return self.doc_abnt.configuracoes.tipo_trabalho == "Artigo Científico"

    def renderizar_cabecalho_artigo(self, doc):
        """Cria os elementos pré-textuais de um artigo na primeira página."""
        p_titulo = doc.add_paragraph()
        p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_titulo = p_titulo.add_run(self.doc_abnt.titulo.upper())
        run_titulo.bold = True
        run_titulo.font.size = self.TAMANHO_FONTE_PADRAO
        p_titulo.paragraph_format.space_after = Pt(18)

        nomes_autores = ', '.join([a.nome_completo for a in self.doc_abnt.autores])
        p_autores = doc.add_paragraph()
        p_autores.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p_autores.add_run(nomes_autores)
        p_autores.paragraph_format.space_after = Pt(18)

        p_resumo_titulo = doc.add_paragraph()
        p_resumo_titulo.add_run("Resumo").bold = True
        p_resumo_titulo.paragraph_format.line_spacing = self.ESPAÇAMENTO_SIMPLES
        
        p_resumo = doc.add_paragraph()
        self.aplicar_estilo_resumo(p_resumo, self.doc_abnt.resumo)
        
        p_kw = doc.add_paragraph()
        run_kw = p_kw.add_run("Palavras-chave: ")
        run_kw.bold = True
        texto_kw = self.doc_abnt.palavras_chave.replace(';', '.') + "."
        p_kw.add_run(texto_kw)
        p_kw.paragraph_format.line_spacing = self.ESPAÇAMENTO_SIMPLES
        p_kw.paragraph_format.space_after = Pt(18)

    def aplicar_estilo_tabela_abnt(self, tabela):
        tbl_pr = tabela._element.xpath('w:tblPr')[0]
        tbl_borders = tbl_pr.find(qn('w:tblBorders'))
        if tbl_borders is None:
            tbl_borders = OxmlElement('w:tblBorders')
            tbl_pr.append(tbl_borders)
        for border in tbl_borders.findall(qn('w:*')):
            tbl_borders.remove(border)
        for border_name in ['top', 'bottom', 'insideH']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single'); border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0'); border.set(qn('w:color'), 'auto')
            tbl_borders.append(border)
        for border_name in ['left', 'right', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tbl_borders.append(border)

    def aplicar_estilo_legenda(self, paragrafo, is_titulo=True):
        paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if is_titulo else WD_PARAGRAPH_ALIGNMENT.LEFT
        paragrafo.paragraph_format.space_before = Pt(0) if is_titulo else Pt(6)
        paragrafo.paragraph_format.space_after = Pt(6) if is_titulo else Pt(12)
        paragrafo.paragraph_format.line_spacing = 1.0
        for run in paragrafo.runs:
            run.font.name = self.FONTE_PADRAO
            run.font.size = self.TAMANHO_FONTE_LEGENDA
            run.font.color.rgb = self.COR_FONTE_PADRAO # CORRIGIDO

    def configurar_pagina_e_estilos(self, doc):
        style = doc.styles['Normal']
        style.font.name = self.FONTE_PADRAO
        style.font.size = self.TAMANHO_FONTE_PADRAO
        style.font.color.rgb = self.COR_FONTE_PADRAO # CORRIGIDO
        style.paragraph_format.line_spacing = self.ESPAÇAMENTO_PADRAO
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        
        # Garante que os estilos de Título também sejam pretos
        for i in range(1, 10):
            style_name = f'Heading {i}'
            if style_name in doc.styles:
                doc.styles[style_name].font.color.rgb = self.COR_FONTE_PADRAO

        citacao_style = doc.styles.add_style('CitacaoLonga', 1)
        citacao_style.base_style = style
        citacao_style.font.size = self.TAMANHO_FONTE_CITACAO_LONGA
        citacao_style.paragraph_format.left_indent = self.RECUO_CITACAO_LONGA
        citacao_style.paragraph_format.line_spacing = self.ESPAÇAMENTO_SIMPLES
        
        ref_style = doc.styles.add_style('Referencias', 1)
        ref_style.base_style = style
        ref_style.paragraph_format.line_spacing = self.ESPAÇAMENTO_SIMPLES
        ref_style.paragraph_format.space_after = Pt(6)
        ref_style.paragraph_format.first_line_indent = 0
        
        for section in doc.sections:
            section.top_margin = self.MARGEM_SUPERIOR
            section.bottom_margin = self.MARGEM_INFERIOR
            section.left_margin = self.MARGEM_ESQUERDA
            section.right_margin = self.MARGEM_DIREITA
            
    def aplicar_estilo_paragrafo_normal(self, paragrafo, texto):
        paragrafo.style = 'Normal'
        paragrafo.paragraph_format.first_line_indent = self.RECUO_PRIMEIRA_LINHA
        paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        run = paragrafo.add_run(texto)
        run.font.name = self.FONTE_PADRAO
        run.font.color.rgb = self.COR_FONTE_PADRAO
        
    def aplicar_estilo_titulo_secao(self, doc, numero, titulo_texto, nivel=1):
        if self.is_artigo:
            titulo_formatado = f"{numero} {titulo_texto}" if numero else titulo_texto.upper()
        else:
            titulo_formatado = f"{numero} {titulo_texto.upper() if nivel == 1 else titulo_texto}" if numero else titulo_texto.upper()
        
        heading = doc.add_heading(titulo_formatado, level=nivel)
        heading.style.font.name = self.FONTE_PADRAO
        heading.style.font.bold = True
        heading.style.font.size = self.TAMANHO_FONTE_PADRAO
        heading.style.font.color.rgb = self.COR_FONTE_PADRAO # CORRIGIDO
        heading.paragraph_format.space_before = Pt(18) if nivel == 1 and not self.is_artigo else Pt(12)
        heading.paragraph_format.space_after = Pt(6)
        heading.paragraph_format.first_line_indent = 0
        
    def aplicar_estilo_natureza_trabalho(self, paragrafo, texto):
        paragrafo.style = 'Normal'
        paragrafo.paragraph_format.left_indent = Cm(8)
        paragrafo.paragraph_format.line_spacing = self.ESPAÇAMENTO_SIMPLES
        paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        paragrafo.add_run(texto)
    
    def aplicar_estilo_resumo(self, paragrafo, texto):
        paragrafo.style = 'Normal'
        paragrafo.paragraph_format.line_spacing = self.ESPAÇAMENTO_SIMPLES
        paragrafo.paragraph_format.first_line_indent = 0
        paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        paragrafo.add_run(texto)
        
    def aplicar_estilo_referencia(self, paragrafo, texto_formatado):
        paragrafo.style = 'Referencias'
        paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        partes = texto_formatado.split('**')
        for i, parte in enumerate(partes):
            run = paragrafo.add_run(parte)
            run.font.color.rgb = self.COR_FONTE_PADRAO # CORRIGIDO
            if i % 2 == 1:
                run.bold = True